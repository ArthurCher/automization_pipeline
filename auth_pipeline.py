# Импорт библиотек для работы с Google Sheets, HTTP-запросами, HTML-парсингом, статистикой и генерацией отчета
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import requests
from bs4 import BeautifulSoup
import re
import statistics
from collections import Counter
import nltk
from nltk.corpus import stopwords
from docx import Document

# Загрузка данных для обработки текста
nltk.download('punkt')
nltk.download('stopwords')

class ContentOptimizer:
    """
    Класс ContentOptimizer отвечает за:
    - Подключение и загрузку данных из Google Sheets
    - Скачивание HTML-страниц по списку URL
    - Извлечение текста и разметки (title, h1, body)
    - Подсчет частоты вхождений ключевых запросов и отдельных слов
    - Удаление выбросов из выборок
    - Сбор и сохранение результатов в виде Word-документа
    """
    def __init__(self, service_account_file, spreadsheet_id, sheet_name):
        # Инициализация подключения к Google Sheets API
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_name(service_account_file, scope)
        self.client = gspread.authorize(creds)
        self.spreadsheet_id = spreadsheet_id
        self.sheet_name = sheet_name

    def fetch_data(self):
        """
        Загружает все строки данных из указанного листа Google Sheets.
        Возвращает список словарей, где каждый словарь содержит данные одной строки.
        """
        sheet = self.client.open_by_key(self.spreadsheet_id).worksheet(self.sheet_name)
        return sheet.get_all_records()

    def download_html(self, url):
        """
        Скачивает HTML-код страницы по URL.
        Если запрос неудачен (ошибка сети, таймаут), возвращает пустую строку.
        """
        try:
            r = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
            r.raise_for_status()
            return r.text
        except:
            return ""

    def extract_text_zones(self, html):
        """
        Принимает HTML-код и извлекает текст из:
        - тега <title>
        - всех тегов <h1>
        - основного текста страницы (body)
        Тексты возвращаются в нижнем регистре.
        """
        soup = BeautifulSoup(html, 'html.parser')
        title = soup.title.string if soup.title else ''
        h1 = ' '.join(h.get_text() for h in soup.find_all('h1'))
        body = soup.get_text()
        return title.lower(), h1.lower(), body.lower()

    def count_word_occurrences(self, text, keywords):
        """
        Считает общее количество вхождений отдельных слов из ключевых запросов.
        Каждое слово из фразы считается отдельно.
        Возвращает целое число.
        """
        counts = 0
        words = set()
        for kw in keywords:
            words.update(kw.lower().split())
        for w in words:
            pattern = r'\b' + re.escape(w) + r'\b'
            counts += len(re.findall(pattern, text))
        return counts

    def remove_outliers(self, values):
        """
        Принимает список числовых значений.
        Возвращает список без выбросов, определенных по методу межквартильного размаха (IQR).
        Если значений меньше 4, возвращает список без изменений.
        """
        if len(values) < 4:
            return values
        q1 = statistics.quantiles(values, n=4)[0]
        q3 = statistics.quantiles(values, n=4)[2]
        iqr = q3 - q1
        return [v for v in values if q1 - 1.5 * iqr <= v <= q3 + 1.5 * iqr]

    def analyze_pages(self, data):
        """
        Основной метод анализа:
        - для каждой строки данных берет список URL страниц (URL serp)
        - скачивает HTML-код страниц
        - извлекает тексты
        - считает:
          * количество вхождений ключевых фраз (по всей фразе)
          * количество вхождений отдельных слов из фраз
          * объем текста
          * LSI-слова (частотные слова, встречающиеся на страницах)
        - возвращает словарь с агрегированными результатами по каждому URL site.
        """
        results = {}
        for row in data:
            site = row['url site']
            urls = eval(row['URL serp (по убыванию частоты)']) if isinstance(row['URL serp (по убыванию частоты)'], str) else row['URL serp (по убыванию частоты)']
            keywords = [row['Фраза']]

            occurrences, word_occurrences, lengths, lsi_words = [], [], [], Counter()
            for url in urls:
                html = self.download_html(url)
                if not html:
                    continue
                title, h1, body = self.extract_text_zones(html)
                text = ' '.join([title, h1, body])
                counts = self.count_keyword_occurrences(text, keywords)
                occurrences.append(sum(counts.values()))
                word_occurrences.append(self.count_word_occurrences(text, keywords))
                lengths.append(len(body))
                tokens = nltk.word_tokenize(text)
                lsi_words.update([t for t in tokens if t not in stopwords.words('russian')])

            results[site] = {
                'mean_occurrences': statistics.mean(self.remove_outliers(occurrences)) if occurrences else 0,
                'mean_word_occurrences': statistics.mean(self.remove_outliers(word_occurrences)) if word_occurrences else 0,
                'mean_length': statistics.mean(self.remove_outliers(lengths)) if lengths else 0,
                'lsi': [w for w, _ in lsi_words.most_common(20)]
            }
        return results

    def generate_report(self, results):
        """
        Формирует Word-документ с итогами анализа.
        Для каждого URL site добавляет:
        - среднее количество вхождений фраз
        - среднее количество вхождений отдельных слов
        - средний размер текста
        - список LSI слов
        Сохраняет файл под названием recommendations.docx.
        """
        doc = Document()
        doc.add_heading('Content Optimization Recommendations', 0)
        for site, stats in results.items():
            doc.add_heading(site, level=1)
            doc.add_paragraph(f"Среднее количество вхождений запросов: {stats['mean_occurrences']}")
            doc.add_paragraph(f"Среднее количество вхождений отдельных слов: {stats['mean_word_occurrences']}")
            doc.add_paragraph(f"Средний объем текста: {stats['mean_length']}")
            doc.add_paragraph("LSI слова: " + ', '.join(stats['lsi']))
        doc.save('recommendations.docx')

if __name__ == '__main__':
    # Запуск всей логики: загрузка данных, анализ страниц, генерация отчета
    optimizer = ContentOptimizer('service_account.json', 'SPREADSHEET_ID', 'CheckTop')
    data = optimizer.fetch_data()
    results = optimizer.analyze_pages(data)
    optimizer.generate_report(results)
